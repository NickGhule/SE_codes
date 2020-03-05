from docx import Document
from docx.shared import Inches
import sys
import os

tempdict = {"TP Rate":0 , "FP Rate" :0, "Precision": 0, "Recall":0 ,"F-Measure": 0,"ROC Area":0,"Class" : 0}
heads = ['Jrip', 'Ridor','J48', 'SC', 'NB', 'IBK' ]
document = Document()


with open("{}.txt".format(sys.argv[1]), "r") as file:
    for line in file:
        count = 0

        if any(head in line for head in heads):
            document.add_paragraph(line)
            pass
        if line == '               TP Rate   FP Rate   Precision   Recall  F-Measure   ROC Area  Class\n':
            table = document.add_table(rows=1, cols=7)
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            for head,i in zip(tempdict.keys(),range(10)):
                hdr_cells[i].text = head

            while(True):
                l = file.readline()
                if not 'Weighted Avg.' in l:
                    temp = l.strip().split()
                    count += 1
                    row_cells = table.add_row().cells
                    for value, row in zip(temp, row_cells):
                        row.text = value
                else:
                    break
            
            var = 'A'
            for i in range(count):
                table.add_column(Inches(1.0))
                hdr_cells = table.rows[0].cells
                hdr_cells[-1].text = var
                var = chr(ord(var)+1)

            for i in range(4):
                l = file.readline()
            for i in range(count):
                l = file.readline()
                data = l.strip().split()
                if data.__len__() == 0:
                    continue
                m = 0
                for j in range(7,7+count):
                    table.cell(i+1,j).text = value = data[m]
                    m += 1
            document.add_paragraph("\n")

try:
    os.mkdir("Converted")
except FileExistsError:
    pass
document.save("Converted\{}.docx".format(sys.argv[1]))